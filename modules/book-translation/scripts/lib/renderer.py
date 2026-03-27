# modules/book-translation/scripts/lib/renderer.py
"""Render markdown content to DOCX with formatting preservation."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm


def parse_markdown_elements(markdown: str) -> list[dict]:
    elements = []
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        # Page break: --- alone on a line
        if re.match(r"^-{3,}$", line):
            elements.append({"type": "page_break"})
            i += 1
            continue
        # ATX headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue
        # Footnote definition: [^N]: text
        footnote_match = re.match(r"^\[\^(\w+)\]:\s+(.+)$", line)
        if footnote_match:
            fn_id = footnote_match.group(1)
            fn_text = footnote_match.group(2).strip()
            elements.append({"type": "footnote_def", "id": fn_id, "text": fn_text})
            i += 1
            continue
        # Blockquote — merge consecutive > lines into single element
        if line.startswith("> "):
            bq_lines = []
            while i < len(lines) and lines[i].rstrip().startswith("> "):
                bq_lines.append(lines[i].rstrip()[2:].strip())
                i += 1
            elements.append({"type": "blockquote", "text": " ".join(bq_lines)})
            continue
        runs = _parse_inline_formatting(line)
        elements.append({"type": "paragraph", "text": line, "runs": runs})
        i += 1
    return elements


def _parse_inline_formatting(text: str) -> list[dict]:
    runs = []
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for match in pattern.finditer(text):
        if match.group(2):
            runs.append({"text": match.group(2), "bold": True, "italic": False})
        elif match.group(3):
            runs.append({"text": match.group(3), "bold": False, "italic": True})
        elif match.group(4):
            runs.append({"text": match.group(4), "bold": False, "italic": False})
    return runs


def render_markdown_to_docx(
    markdown: str,
    output_path: Path,
    font_name: str = "Times New Roman",
    font_size: int = 12,
    heading1_size: int = 16,
    heading2_size: int = 14,
    heading3_size: int = 13,
    margin_top: float = 2.5,
    margin_bottom: float = 2.5,
    margin_left: float = 3.0,
    margin_right: float = 2.0,
    line_spacing: float = 1.5,
) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(margin_top)
        section.bottom_margin = Cm(margin_bottom)
        section.left_margin = Cm(margin_left)
        section.right_margin = Cm(margin_right)

    heading_sizes = {1: heading1_size, 2: heading2_size, 3: heading3_size}
    elements = parse_markdown_elements(markdown)

    for elem in elements:
        if elem["type"] == "page_break":
            doc.add_page_break()
        elif elem["type"] == "heading":
            level = min(elem["level"], 3)
            para = doc.add_heading(elem["text"], level=level)
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(heading_sizes.get(level, heading3_size))
        elif elem["type"] == "blockquote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            run = para.add_run(elem["text"])
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.italic = True
        elif elem["type"] == "footnote_def":
            # Render footnote definition as smaller italic text
            para = doc.add_paragraph()
            run = para.add_run(f"[^{elem['id']}]: {elem['text']}")
            run.font.name = font_name
            run.font.size = Pt(max(font_size - 2, 8))
            run.italic = True
        elif elem["type"] == "paragraph":
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = line_spacing
            for run_data in elem.get("runs", []):
                run = para.add_run(run_data["text"])
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = run_data.get("bold", False)
                run.italic = run_data.get("italic", False)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
